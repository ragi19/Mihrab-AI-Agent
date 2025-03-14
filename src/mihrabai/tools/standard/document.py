"""
Document processing tools for working with various file formats
"""

import base64
import json
import mimetypes
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from ...utils.logging import get_logger
from ..base import BaseTool

logger = get_logger("tools.document")


class DocumentExtractorTool(BaseTool):
    """Tool for extracting text content from various document formats"""

    def __init__(self):
        super().__init__(
            name="document_extractor",
            description="Extract text content from various document formats",
        )
        self._parameters = {
            "file_path": {"type": "string", "description": "Path to the document file"},
            "format": {
                "type": "string",
                "description": "Format of the document (auto-detected if not specified)",
                "enum": ["txt", "markdown", "html", "pdf", "docx", "auto"],
                "default": "auto",
            },
            "start_page": {
                "type": "integer",
                "description": "Starting page for extraction (for multi-page documents)",
                "default": 1,
            },
            "end_page": {
                "type": "integer",
                "description": "Ending page for extraction (for multi-page documents)",
                "default": 0,
            },
            "include_metadata": {
                "type": "boolean",
                "description": "Include document metadata in the output",
                "default": True,
            },
        }
        self._required_params = ["file_path"]
        logger.info("Initialized tool: document_extractor")

    def _get_parameters_schema(self) -> Dict[str, Any]:
        """Get the JSON schema for tool parameters"""
        return {
            "type": "object",
            "properties": self._parameters,
            "required": self._required_params,
        }

    async def _execute(
        self,
        file_path: str,
        format: str = "auto",
        start_page: int = 1,
        end_page: int = 0,
        include_metadata: bool = True,
        **kwargs,
    ) -> Dict[str, Any]:
        """Execute the document extractor tool

        Args:
            file_path: Path to the document file
            format: Format of the document
            start_page: Starting page for extraction
            end_page: Ending page for extraction
            include_metadata: Include document metadata in the output

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}

            # Auto-detect format if not specified
            if format == "auto":
                format = self._detect_format(file_path)

            # Extract text based on format
            if format == "txt":
                text, metadata = self._extract_from_txt(file_path)
            elif format == "markdown":
                text, metadata = self._extract_from_markdown(file_path)
            elif format == "html":
                text, metadata = self._extract_from_html(file_path)
            elif format == "pdf":
                text, metadata = self._extract_from_pdf(file_path, start_page, end_page)
            elif format == "docx":
                text, metadata = self._extract_from_docx(file_path)
            else:
                return {"error": f"Unsupported format: {format}"}

            # Prepare result
            result = {
                "text": text,
                "format": format,
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "character_count": len(text),
                "word_count": len(re.findall(r"\b\w+\b", text)),
            }

            # Include metadata if requested
            if include_metadata and metadata:
                result["metadata"] = metadata

            return result
        except Exception as e:
            logger.error(f"Document extraction error: {e}")
            return {
                "error": f"Document extraction error: {str(e)}",
                "file_path": file_path,
                "format": format,
            }

    def _detect_format(self, file_path: str) -> str:
        """Detect document format from file extension

        Args:
            file_path: Path to the document file

        Returns:
            Detected format
        """
        extension = os.path.splitext(file_path)[1].lower()

        if extension == ".txt":
            return "txt"
        elif extension in [".md", ".markdown"]:
            return "markdown"
        elif extension in [".html", ".htm"]:
            return "html"
        elif extension == ".pdf":
            return "pdf"
        elif extension in [".docx", ".doc"]:
            return "docx"
        else:
            # Default to txt for unknown formats
            return "txt"

    def _extract_from_txt(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from a plain text file

        Args:
            file_path: Path to the text file

        Returns:
            Tuple of (extracted text, metadata)
        """
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            metadata = {"encoding": "utf-8", "line_count": text.count("\n") + 1}

            return text, metadata
        except UnicodeDecodeError:
            # Try with different encodings
            encodings = ["latin-1", "cp1252", "iso-8859-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()

                    metadata = {
                        "encoding": encoding,
                        "line_count": text.count("\n") + 1,
                    }

                    return text, metadata
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, raise the original error
            raise

    def _extract_from_markdown(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from a Markdown file

        Args:
            file_path: Path to the Markdown file

        Returns:
            Tuple of (extracted text, metadata)
        """
        # Extract text (same as txt)
        text, basic_metadata = self._extract_from_txt(file_path)

        # Extract additional Markdown-specific metadata
        headers = re.findall(r"^#+\s+(.+)$", text, re.MULTILINE)
        links = re.findall(r"\[([^\]]+)\]\(([^)]+)\)", text)
        code_blocks = re.findall(r"```(?:\w+)?\n(.*?)```", text, re.DOTALL)

        metadata = {
            **basic_metadata,
            "headers": headers,
            "link_count": len(links),
            "code_block_count": len(code_blocks),
        }

        return text, metadata

    def _extract_from_html(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from an HTML file

        Args:
            file_path: Path to the HTML file

        Returns:
            Tuple of (extracted text, metadata)
        """
        try:
            # Try to import BeautifulSoup
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()

            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # Extract text
            text = soup.get_text(separator=" ", strip=True)

            # Extract metadata
            title = soup.title.string if soup.title else None
            meta_tags = {
                meta.get("name", meta.get("property", "unknown")): meta.get("content")
                for meta in soup.find_all("meta")
                if meta.get("content")
            }

            metadata = {
                "title": title,
                "meta_tags": meta_tags,
                "links": [a.get("href") for a in soup.find_all("a") if a.get("href")],
                "images": [
                    img.get("src") for img in soup.find_all("img") if img.get("src")
                ],
                "headings": {
                    f"h{i}": [h.get_text(strip=True) for h in soup.find_all(f"h{i}")]
                    for i in range(1, 7)
                },
            }

            return text, metadata
        except ImportError:
            # Fallback if BeautifulSoup is not available
            logger.warning(
                "BeautifulSoup not available, using regex-based HTML extraction"
            )

            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()

            # Simple regex-based extraction
            text = re.sub(r"<[^>]+>", " ", html_content)
            text = re.sub(r"\s+", " ", text).strip()

            # Extract basic metadata with regex
            title_match = re.search(r"<title>([^<]+)</title>", html_content)
            title = title_match.group(1) if title_match else None

            metadata = {
                "title": title,
                "note": "Limited metadata extracted (BeautifulSoup not available)",
            }

            return text, metadata

    def _extract_from_pdf(
        self, file_path: str, start_page: int = 1, end_page: int = 0
    ) -> tuple[str, Dict[str, Any]]:
        """Extract text from a PDF file

        Args:
            file_path: Path to the PDF file
            start_page: Starting page for extraction
            end_page: Ending page for extraction (0 for all pages)

        Returns:
            Tuple of (extracted text, metadata)
        """
        try:
            # Try to import PyPDF2
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            num_pages = len(reader.pages)

            # Adjust page range
            if start_page < 1:
                start_page = 1

            if end_page <= 0 or end_page > num_pages:
                end_page = num_pages

            # Extract text from specified pages
            text = ""
            for page_num in range(start_page - 1, end_page):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"

            # Extract metadata
            metadata = {
                "total_pages": num_pages,
                "extracted_pages": list(range(start_page, end_page + 1)),
                "pdf_info": reader.metadata if hasattr(reader, "metadata") else {},
            }

            return text, metadata
        except ImportError:
            # Fallback message if PyPDF2 is not available
            logger.error("PyPDF2 not available for PDF extraction")
            raise ImportError(
                "PyPDF2 library is required for PDF extraction. Please install it with 'pip install PyPDF2'."
            )

    def _extract_from_docx(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from a DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted text, metadata)
        """
        try:
            # Try to import python-docx
            import docx

            doc = docx.Document(file_path)

            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Extract metadata
            metadata = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": (
                    doc.core_properties.created.isoformat()
                    if doc.core_properties.created
                    else None
                ),
                "modified": (
                    doc.core_properties.modified.isoformat()
                    if doc.core_properties.modified
                    else None
                ),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            return text, metadata
        except ImportError:
            # Fallback message if python-docx is not available
            logger.error("python-docx not available for DOCX extraction")
            raise ImportError(
                "python-docx library is required for DOCX extraction. Please install it with 'pip install python-docx'."
            )


class DocumentConverterTool(BaseTool):
    """Tool for converting documents between different formats"""

    def __init__(self):
        super().__init__(
            name="document_converter",
            description="Convert documents between different formats",
        )
        self._parameters = {
            "input_file": {
                "type": "string",
                "description": "Path to the input document file",
            },
            "output_format": {
                "type": "string",
                "description": "Target format for conversion",
                "enum": ["txt", "markdown", "html", "pdf", "docx"],
                "default": "txt",
            },
            "output_file": {
                "type": "string",
                "description": "Path for the output file (optional, generated if not provided)",
            },
            "options": {
                "type": "object",
                "description": "Additional options for conversion",
                "default": {},
            },
        }
        self._required_params = ["input_file", "output_format"]
        logger.info("Initialized tool: document_converter")

    def _get_parameters_schema(self) -> Dict[str, Any]:
        """Get the JSON schema for tool parameters"""
        return {
            "type": "object",
            "properties": self._parameters,
            "required": self._required_params,
        }

    async def _execute(
        self,
        input_file: str,
        output_format: str,
        output_file: Optional[str] = None,
        options: Dict[str, Any] = None,
        **kwargs,
    ) -> Dict[str, Any]:
        """Execute the document converter tool

        Args:
            input_file: Path to the input document file
            output_format: Target format for conversion
            output_file: Path for the output file
            options: Additional options for conversion

        Returns:
            Dictionary with conversion result
        """
        try:
            # Check if input file exists
            if not os.path.exists(input_file):
                return {"error": f"Input file not found: {input_file}"}

            # Generate output file path if not provided
            if not output_file:
                input_base = os.path.splitext(input_file)[0]
                output_file = f"{input_base}.{output_format}"

            # Detect input format
            input_format = os.path.splitext(input_file)[1].lower().lstrip(".")
            if input_format in ["doc", "docx"]:
                input_format = "docx"

            # Check if input and output formats are the same
            if input_format == output_format:
                return {
                    "error": f"Input and output formats are the same: {input_format}"
                }

            # Initialize options
            if options is None:
                options = {}

            # Perform conversion
            if self._is_pandoc_conversion_needed(input_format, output_format):
                success, message = self._convert_with_pandoc(
                    input_file, output_file, input_format, output_format, options
                )
            else:
                success, message = self._convert_with_custom_logic(
                    input_file, output_file, input_format, output_format, options
                )

            if not success:
                return {"error": message}

            return {
                "status": "success",
                "input_file": input_file,
                "output_file": output_file,
                "input_format": input_format,
                "output_format": output_format,
                "message": message,
            }
        except Exception as e:
            logger.error(f"Document conversion error: {e}")
            return {
                "error": f"Document conversion error: {str(e)}",
                "input_file": input_file,
                "output_format": output_format,
            }

    def _is_pandoc_conversion_needed(
        self, input_format: str, output_format: str
    ) -> bool:
        """Check if pandoc is needed for this conversion

        Args:
            input_format: Input document format
            output_format: Output document format

        Returns:
            True if pandoc is needed, False otherwise
        """
        # Complex conversions that require pandoc
        complex_formats = ["pdf", "docx", "html"]

        if input_format in complex_formats or output_format in complex_formats:
            return True

        return False

    def _convert_with_pandoc(
        self,
        input_file: str,
        output_file: str,
        input_format: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> tuple[bool, str]:
        """Convert document using pandoc

        Args:
            input_file: Path to the input document file
            output_file: Path for the output file
            input_format: Input document format
            output_format: Output document format
            options: Additional options for conversion

        Returns:
            Tuple of (success, message)
        """
        try:
            import subprocess

            # Build pandoc command
            cmd = ["pandoc", input_file, "-o", output_file]

            # Add format-specific options
            if input_format == "pdf" and output_format != "pdf":
                cmd.append("--pdf-engine=pdftotext")

            # Add any additional options
            for key, value in options.items():
                if isinstance(value, bool) and value:
                    cmd.append(f"--{key}")
                else:
                    cmd.append(f"--{key}={value}")

            # Execute pandoc
            result = subprocess.run(cmd, capture_output=True, text=True, check=False)

            if result.returncode != 0:
                return False, f"Pandoc conversion failed: {result.stderr}"

            return True, "Conversion completed successfully with pandoc"
        except FileNotFoundError:
            return (
                False,
                "Pandoc not found. Please install pandoc: https://pandoc.org/installing.html",
            )
        except Exception as e:
            return False, f"Pandoc conversion error: {str(e)}"

    def _convert_with_custom_logic(
        self,
        input_file: str,
        output_file: str,
        input_format: str,
        output_format: str,
        options: Dict[str, Any],
    ) -> tuple[bool, str]:
        """Convert document using custom logic for simple conversions

        Args:
            input_file: Path to the input document file
            output_file: Path for the output file
            input_format: Input document format
            output_format: Output document format
            options: Additional options for conversion

        Returns:
            Tuple of (success, message)
        """
        try:
            # Simple text-based conversions
            if input_format in ["txt", "markdown"] and output_format in [
                "txt",
                "markdown",
            ]:
                with open(input_file, "r", encoding="utf-8") as infile:
                    content = infile.read()

                # Convert Markdown to plain text (simple stripping of markdown syntax)
                if input_format == "markdown" and output_format == "txt":
                    # Remove headers
                    content = re.sub(r"^#+\s+", "", content, flags=re.MULTILINE)
                    # Remove emphasis
                    content = re.sub(r"\*\*(.*?)\*\*", r"\1", content)
                    content = re.sub(r"\*(.*?)\*", r"\1", content)
                    # Remove links
                    content = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", content)
                    # Remove code blocks
                    content = re.sub(r"```.*?```", "", content, flags=re.DOTALL)
                    # Remove inline code
                    content = re.sub(r"`(.*?)`", r"\1", content)

                # Convert plain text to Markdown (just copy)
                # No special processing needed for txt to markdown

                with open(output_file, "w", encoding="utf-8") as outfile:
                    outfile.write(content)

                return True, f"Converted {input_format} to {output_format} successfully"
            else:
                return (
                    False,
                    f"Custom conversion from {input_format} to {output_format} not supported",
                )
        except Exception as e:
            return False, f"Custom conversion error: {str(e)}"


class DocumentSplitterTool(BaseTool):
    """Tool for splitting documents into smaller chunks"""

    def __init__(self):
        super().__init__(
            name="document_splitter",
            description="Split documents into smaller chunks for processing",
        )
        self._parameters = {
            "text": {"type": "string", "description": "Text content to split"},
            "split_method": {
                "type": "string",
                "description": "Method to use for splitting",
                "enum": ["chunk", "sentence", "paragraph", "heading"],
                "default": "chunk",
            },
            "chunk_size": {
                "type": "integer",
                "description": "Size of each chunk in characters (for chunk method)",
                "default": 1000,
            },
            "chunk_overlap": {
                "type": "integer",
                "description": "Overlap between chunks in characters (for chunk method)",
                "default": 100,
            },
            "separator": {
                "type": "string",
                "description": "Custom separator for splitting (for sentence/paragraph methods)",
                "default": "",
            },
        }
        self._required_params = ["text"]
        logger.info("Initialized tool: document_splitter")

    def _get_parameters_schema(self) -> Dict[str, Any]:
        """Get the JSON schema for tool parameters"""
        return {
            "type": "object",
            "properties": self._parameters,
            "required": self._required_params,
        }

    async def _execute(
        self,
        text: str,
        split_method: str = "chunk",
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
        separator: str = "",
        **kwargs,
    ) -> Dict[str, Any]:
        """Execute the document splitter tool

        Args:
            text: Text content to split
            split_method: Method to use for splitting
            chunk_size: Size of each chunk in characters
            chunk_overlap: Overlap between chunks in characters
            separator: Custom separator for splitting

        Returns:
            Dictionary with split chunks
        """
        try:
            # Validate parameters
            if chunk_size <= 0:
                return {"error": "chunk_size must be greater than 0"}

            if chunk_overlap < 0 or chunk_overlap >= chunk_size:
                return {"error": "chunk_overlap must be between 0 and chunk_size"}

            # Split text based on method
            if split_method == "chunk":
                chunks = self._split_by_chunk(text, chunk_size, chunk_overlap)
            elif split_method == "sentence":
                chunks = self._split_by_sentence(text, separator)
            elif split_method == "paragraph":
                chunks = self._split_by_paragraph(text, separator)
            elif split_method == "heading":
                chunks = self._split_by_heading(text)
            else:
                return {"error": f"Invalid split method: {split_method}"}

            # Calculate statistics
            chunk_lengths = [len(chunk) for chunk in chunks]

            return {
                "chunks": chunks,
                "chunk_count": len(chunks),
                "avg_chunk_length": sum(chunk_lengths) / max(1, len(chunks)),
                "min_chunk_length": min(chunk_lengths) if chunks else 0,
                "max_chunk_length": max(chunk_lengths) if chunks else 0,
                "total_length": len(text),
                "split_method": split_method,
            }
        except Exception as e:
            logger.error(f"Document splitting error: {e}")
            return {
                "error": f"Document splitting error: {str(e)}",
                "split_method": split_method,
            }

    def _split_by_chunk(
        self, text: str, chunk_size: int, chunk_overlap: int
    ) -> List[str]:
        """Split text into overlapping chunks of specified size

        Args:
            text: Text to split
            chunk_size: Size of each chunk in characters
            chunk_overlap: Overlap between chunks in characters

        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            # Calculate end position
            end = start + chunk_size

            # Adjust end to avoid splitting words
            if end < text_length:
                # Look for whitespace to break at
                while end > start and not text[end].isspace():
                    end -= 1

                # If no whitespace found, just use the chunk size
                if end == start:
                    end = start + chunk_size
            else:
                end = text_length

            # Add chunk to list
            chunks.append(text[start:end])

            # Calculate next start position with overlap
            start = end - chunk_overlap

            # Ensure we make progress
            if start >= end:
                start = end

        return chunks

    def _split_by_sentence(self, text: str, separator: str = "") -> List[str]:
        """Split text into sentences

        Args:
            text: Text to split
            separator: Custom separator (if empty, use default sentence boundaries)

        Returns:
            List of sentences
        """
        if separator:
            # Split by custom separator
            return [s.strip() for s in text.split(separator) if s.strip()]
        else:
            # Split by sentence boundaries
            sentence_endings = r"(?<=[.!?])\s+"
            sentences = re.split(sentence_endings, text)
            return [s.strip() for s in sentences if s.strip()]

    def _split_by_paragraph(self, text: str, separator: str = "") -> List[str]:
        """Split text into paragraphs

        Args:
            text: Text to split
            separator: Custom separator (if empty, use default paragraph boundaries)

        Returns:
            List of paragraphs
        """
        if separator:
            # Split by custom separator
            return [p.strip() for p in text.split(separator) if p.strip()]
        else:
            # Split by double newlines (common paragraph separator)
            paragraphs = re.split(r"\n\s*\n", text)
            return [p.strip() for p in paragraphs if p.strip()]

    def _split_by_heading(self, text: str) -> List[str]:
        """Split text by headings (Markdown style)

        Args:
            text: Text to split (assumed to be in Markdown format)

        Returns:
            List of sections (heading + content)
        """
        # Find all headings
        heading_pattern = r"^(#{1,6})\s+(.+)$"
        matches = list(re.finditer(heading_pattern, text, re.MULTILINE))

        if not matches:
            # No headings found, return the whole text
            return [text]

        chunks = []

        # Process each section
        for i, match in enumerate(matches):
            # Get the start of this section
            start_pos = match.start()

            # Get the end of this section (start of next section or end of text)
            end_pos = matches[i + 1].start() if i < len(matches) - 1 else len(text)

            # Extract the section
            section = text[start_pos:end_pos].strip()
            if section:
                chunks.append(section)

        # Check if there's content before the first heading
        if matches and matches[0].start() > 0:
            first_chunk = text[: matches[0].start()].strip()
            if first_chunk:
                chunks.insert(0, first_chunk)

        return chunks
